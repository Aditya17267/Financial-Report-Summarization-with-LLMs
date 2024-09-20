from docx import Document
import logging

logging.basicConfig(level=logging.INFO)

def save_text_to_docx(text_data, output_path):
    """
    Saves text data to a Word document.
    
    Args:
        text_data (str): A string containing the entire text data.
        output_path (str): Path to the output Word document.
    """
    try:
        doc = Document()
        lines = text_data.split('\n')

        for line in lines:
            if line.startswith('####'):
                doc.add_heading(line.replace('####', '').strip(), level=1)
            elif line.startswith('###'):
                doc.add_heading(line.replace('###', '').strip(), level=2)
            elif line.startswith('- **'):
                bold_text, normal_text = line.split('**')[1], line.split('**')[2].strip()
                p = doc.add_paragraph()
                p.add_run(bold_text + ':').bold = True
                p.add_run(' ' + normal_text)
            elif line.strip() == '':
                continue
            else:
                doc.add_paragraph(line.strip())

        doc.save(output_path)
    except IOError:
        logging.error(f"IO error while saving text to DOCX at {output_path}.")
    except Exception as e:
        logging.error(f"Unexpected error saving text to DOCX: {e}")
